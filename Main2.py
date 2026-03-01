import random
from funcs import *
import csv
import numpy as np



# Requirements
# Takes 16 players (potential for changing this number in the future for top 8 and top 4)
# Each player plays every other player
# A tracking variable that shows difficulty level for each player
# A variable that can control how equal difficulties are for players (Either have this as low as possible (least variation), or equalize scores at the end)

# Inputs: player_count, difficulty_variation, pattern_boards? (boolean), board and dice ranges.
# Outputs: 15(+) rounds of boards, dice rolls, and calculated difficulty per player ( or expected score ).

# For 16 players, 15 rounds, 3 boards, 5 rolls/board. For 8, 7 rounds, 2 boards, 3-4 rolls/board. For 4, 4 rounds, 1 board, 4 rolls.
# Foe 32 players, 31 rounds, 6 boards, 5(+1/6) rolls/board

player_count = 16

rounds_arr = {32:31, 16:15, 8:7, 4:4}
boards_arr = {32:6, 16:3, 8:2, 4:1}

round_count = rounds_arr[player_count]
board_count = boards_arr[player_count]
dice_count = round(round_count/board_count - 0.5)

bug_check = False

# This increases if the amount of dice rolls on the last board of the round is increased (example is for 8 players on the second board it has 4 instead of 3 so this value would be 1)
extra_dice = 0

if round_count % board_count != 0:
    extra_dice += (round_count) - (dice_count * board_count)

# Bug Check
if bug_check:
    print("board_count: " + str(board_count))
    print("dice_count: " + str(dice_count))
    print("extra_dice: " + str(extra_dice))

# Create ranges for boards
board_ranges = []
for i in range(board_count):
    board_ranges.append(round(random.random() * 550 + 72))

# Bug Check
if bug_check:
    print("board_ranges : " + str(board_ranges))

# Generate boards
boards = []
include_pattern_boards = True

# If we include pattern boards, it should be a random amount from 0-50% of the total. If there is more than 2 boards, lets require one to be a pattern.
pattern_board_count = round(random.random()*0.49 * board_count)

if board_count > 2 and pattern_board_count < 1:
    pattern_board_count = 1

for i in range(board_count - pattern_board_count):
    boards.append(generate_random_board(board_ranges[i]))

for i in range(pattern_board_count):
    multiple_list = []
    multiple_count = round(random.random() * 3 - 0.51) + 1

    # This is the multiple, for double and triple multiples, this should be the average increase.
    multiple = round(board_ranges[i +  (board_count - pattern_board_count)] / 36)

    # Starting number
    starting_number = round(random.random() * 20)

    # Single number pattern 
    if multiple_count == 1:
        multiple_list.append(multiple)
        boards.append(generate_pattern_board(multiple=multiple_list, startingNumber=starting_number))
    

    # Double number pattern
    if multiple_count == 2:
        multiple_list.append(round((random.random() * 2 - 0.5) * multiple*2))
        multiple_list.append(multiple*2 - multiple_list[0])
        boards.append(generate_pattern_board(multiple=multiple_list, startingNumber=starting_number))
    
    # Triple number pattern
    if multiple_count == 3:
        # Create three different multiples
        multiple_list.append(round((random.random() * 2 - 0.5) * multiple*3))
        multiple_list.append(round((random.random() * 2 - 0.5) * multiple*3))
        # Ensure the third multiple makes the sum equal to the original multiple
        multiple_list.append(multiple*3 - sum(multiple_list))

        # Ensure there are no negative numbers. (We check the largest amount that the multiple will subtract, then add that to the starting number)
        most_negative = 0
        for i in multiple_list:           
            if i < 0:
                most_negative += i   
        starting_number -= most_negative

        boards.append(generate_pattern_board(multiple=multiple_list, startingNumber=starting_number))

if bug_check:
    print("board_count: "+ str(board_count))
    print("pattern_board_count: " + str(pattern_board_count))
    print("multiple count: " + str(multiple_count))
    print("multiple_list: " + str(multiple_list))
    print("boards: ")
    for i in range(len(boards)):
        print_board(boards[i])
    
# Keeping track of difficulty for all numbers, all rolls is gonna be a challenge.
# We'll start with each board being represented in board_difficulties
# board_difficulties[0] is the first board and so on
# board_difficulties[0][0] will be the expected score on the first board with the first dice

# We have 2 dice lists, a normal one and an extensive one.  The normal one will use less compute but be more comprehensive
# All dice are reduced to smallest power, 4->2, 9->3 etc.

normal_dice_list = [[2, 2, 2], [2, 2, 3], [2, 2, 5], [2, 2, 6], [2, 2, 7], [2, 2, 10], [2, 3, 3], [2, 3, 5], [2, 3, 6],
            [2, 3, 7], [2, 3, 10], [2, 3, 11], [2, 3, 12], [2, 3, 13], [2, 3, 14], [2, 3, 15], [2, 3, 17],
            [2, 3, 18], [2, 3, 19], [2, 3, 20], [2, 5, 5], [2, 5, 6], [2, 5, 10], [2, 6, 6], [2, 6, 7], [2, 6, 12],
            [3, 3, 5], [3, 3, 6], [3, 3, 12], [3, 5, 5], [3, 5, 6], [3, 5, 7], [3, 5, 15], [3, 6, 6],
            [3, 6, 7], [3, 6, 10], [3, 6, 12], [3, 6, 18]]

extensive_dice_list = [[1, 2, 2], [1, 2, 3], [1, 2, 5], [1, 2, 6], [1, 2, 7], [1, 2, 10], [1, 2, 11], [1, 2, 12], [1, 2, 13], [1, 2, 14], [1, 2, 15], [1, 2, 17], [1, 2, 18], [1, 2, 19], [1, 2, 20], [1, 3, 3], [1, 3, 5], [1, 3, 6], [1, 3, 7], [1, 3, 10], [1, 3, 11], [1, 3, 12], [1, 3, 13], [1, 3, 14], [1, 3, 15], [1, 3, 17], [1, 3, 18], [1, 3, 19], [1, 3, 20], [1, 5, 5], [1, 5, 6], [1, 5, 7], [1, 5, 10], [1, 5, 11], [1, 5, 12], [1, 5, 13], [1, 5, 14], [1, 5, 15], [1, 5, 17], [1, 5, 18], [1, 5, 19], [1, 5, 20], [1, 6, 6], [1, 6, 7], [1, 6, 10], [1, 6, 11], [1, 6, 12], [1, 6, 13], [1, 6, 14], [1, 6, 15], [1, 6, 17], [1, 6, 18], [1, 6, 19], [1, 6, 20], [1, 7, 7], [1, 7, 10], [1, 7, 11], [1, 7, 12], [1, 7, 13], [1, 7, 14], [1, 7, 15], [1, 7, 17], [1, 7, 18], [1, 7, 19], [1, 7, 20], [1, 10, 10], [1, 10, 11], [1, 10, 12], [1, 10, 13], [1, 10, 14], [1, 10, 15], [1, 10, 17], [1, 10, 18], [1, 10, 19], [1, 10, 20], [1, 
11, 11], [1, 11, 12], [1, 11, 13], [1, 11, 14], [1, 11, 15], [1, 11, 17], [1, 11, 18], [1, 11, 19], [1, 11, 20], [1, 12, 12], [1, 12, 13], [1, 12, 14], [1, 12, 15], [1, 12, 17], [1, 12, 18], [1, 12, 19], [1, 12, 20], [1, 13, 
13], [1, 13, 14], [1, 13, 15], [1, 13, 17], [1, 13, 18], [1, 13, 19], [1, 13, 20], [1, 14, 14], [1, 14, 15], [1, 14, 17], [1, 14, 18], [1, 14, 19], [1, 14, 20], [1, 15, 15], [1, 15, 17], [1, 15, 18], [1, 15, 19], [1, 15, 20], [1, 17, 17], [1, 17, 18], [1, 17, 19], [1, 17, 20], [1, 18, 18], [1, 18, 19], [1, 18, 20], [1, 19, 19], [1, 19, 20], [1, 20, 20], [2, 2, 2], [2, 2, 3], [2, 2, 5], [2, 2, 6], [2, 2, 7], [2, 2, 10], [2, 2, 11], [2, 2, 12], [2, 2, 13], [2, 2, 14], [2, 2, 15], [2, 2, 17], [2, 2, 18], [2, 2, 19], [2, 2, 20], [2, 3, 3], [2, 3, 5], [2, 3, 6], [2, 3, 7], [2, 3, 10], [2, 3, 11], [2, 3, 12], [2, 3, 13], [2, 3, 14], [2, 3, 15], [2, 3, 17], [2, 3, 18], [2, 
3, 19], [2, 3, 20], [2, 5, 5], [2, 5, 6], [2, 5, 7], [2, 5, 10], [2, 5, 11], [2, 5, 12], [2, 5, 13], [2, 5, 14], [2, 5, 15], [2, 5, 17], [2, 5, 18], [2, 5, 19], [2, 5, 20], [2, 6, 6], [2, 6, 7], [2, 6, 10], [2, 6, 11], [2, 6, 12], [2, 6, 13], [2, 6, 14], [2, 6, 15], [2, 6, 17], [2, 6, 18], [2, 6, 19], [2, 6, 20], [2, 7, 7], [2, 7, 10], [2, 7, 11], [2, 7, 12], [2, 7, 13], [2, 7, 14], [2, 7, 15], [2, 7, 17], [2, 7, 18], [2, 7, 19], [2, 7, 20], [2, 
10, 10], [2, 10, 11], [2, 10, 12], [2, 10, 13], [2, 10, 14], [2, 10, 15], [2, 10, 17], [2, 10, 18], [2, 10, 19], [2, 10, 20], [2, 11, 11], [2, 11, 12], [2, 11, 13], [2, 11, 14], [2, 11, 15], [2, 11, 17], [2, 11, 18], [2, 11, 
19], [2, 11, 20], [2, 12, 12], [2, 12, 13], [2, 12, 14], [2, 12, 15], [2, 12, 17], [2, 12, 18], [2, 12, 19], [2, 12, 20], [2, 13, 13], [2, 13, 14], [2, 13, 15], [2, 13, 17], [2, 13, 18], [2, 13, 19], [2, 13, 20], [2, 14, 14], [2, 14, 15], [2, 14, 17], [2, 14, 18], [2, 14, 19], [2, 14, 20], [2, 15, 15], [2, 15, 17], [2, 15, 18], [2, 15, 19], [2, 15, 20], [2, 17, 17], [2, 17, 18], [2, 17, 19], [2, 17, 20], [2, 18, 18], [2, 18, 19], [2, 18, 20], [2, 19, 19], [2, 19, 20], [2, 20, 20], [3, 3, 3], [3, 3, 5], [3, 3, 6], [3, 3, 7], [3, 3, 10], [3, 3, 11], [3, 3, 12], [3, 3, 13], [3, 3, 14], [3, 3, 15], [3, 3, 17], [3, 3, 18], [3, 3, 19], [3, 3, 20], [3, 5, 5], [3, 5, 6], [3, 5, 7], [3, 5, 10], [3, 5, 11], [3, 5, 12], [3, 5, 13], [3, 5, 14], [3, 5, 15], [3, 5, 17], [3, 5, 18], [3, 5, 19], [3, 5, 20], [3, 6, 6], [3, 6, 7], [3, 6, 10], [3, 6, 11], [3, 6, 12], [3, 6, 13], [3, 6, 14], [3, 6, 15], [3, 6, 17], [3, 6, 18], [3, 6, 19], [3, 6, 20], [3, 7, 7], [3, 7, 10], [3, 7, 11], [3, 7, 12], [3, 7, 13], [3, 7, 14], [3, 7, 15], [3, 7, 17], [3, 7, 18], [3, 7, 19], [3, 7, 20], [3, 10, 10], [3, 10, 11], [3, 10, 12], [3, 10, 13], [3, 10, 14], [3, 10, 15], [3, 10, 17], [3, 10, 18], [3, 10, 19], [3, 10, 20], [3, 11, 11], [3, 11, 12], [3, 11, 13], [3, 11, 14], [3, 11, 15], [3, 11, 17], [3, 11, 18], [3, 11, 19], [3, 11, 20], [3, 12, 12], [3, 12, 13], [3, 12, 14], [3, 12, 15], [3, 12, 17], [3, 12, 18], [3, 12, 19], [3, 12, 20], [3, 13, 13], [3, 13, 14], [3, 13, 15], [3, 13, 17], [3, 13, 18], [3, 13, 19], [3, 13, 20], [3, 14, 14], [3, 14, 15], [3, 14, 17], [3, 14, 18], [3, 14, 19], [3, 14, 20], [3, 15, 15], [3, 15, 17], [3, 15, 18], [3, 15, 19], [3, 15, 20], [3, 17, 17], [3, 17, 18], [3, 17, 19], [3, 17, 20], [3, 18, 18], [3, 18, 19], [3, 18, 20], [3, 19, 19], [3, 19, 20], [3, 20, 20], [5, 5, 5], [5, 5, 6], [5, 5, 7], [5, 5, 10], [5, 5, 11], [5, 5, 12], [5, 5, 13], [5, 5, 14], [5, 5, 15], [5, 5, 17], [5, 5, 18], [5, 5, 19], [5, 5, 20], [5, 6, 6], [5, 6, 7], [5, 6, 10], [5, 6, 11], [5, 6, 12], [5, 6, 13], [5, 6, 14], [5, 6, 15], [5, 6, 17], [5, 6, 18], [5, 6, 19], [5, 6, 20], [5, 7, 7], [5, 7, 10], [5, 7, 11], [5, 7, 12], [5, 7, 13], [5, 7, 14], [5, 7, 15], [5, 7, 17], [5, 7, 18], [5, 7, 19], [5, 7, 20], [5, 10, 10], [5, 10, 11], [5, 10, 12], [5, 10, 13], [5, 10, 14], [5, 10, 15], [5, 10, 17], [5, 10, 18], [5, 10, 19], [5, 10, 20], [5, 11, 11], [5, 11, 12], [5, 11, 13], [5, 11, 14], [5, 11, 15], [5, 11, 17], [5, 11, 18], [5, 11, 19], [5, 11, 20], [5, 12, 12], [5, 12, 13], [5, 12, 14], [5, 12, 15], [5, 12, 17], [5, 12, 18], [5, 12, 19], [5, 12, 20], [5, 13, 13], [5, 13, 14], [5, 13, 15], [5, 13, 17], [5, 13, 18], [5, 13, 19], [5, 13, 20], [5, 14, 14], [5, 14, 15], [5, 14, 17], [5, 14, 18], [5, 14, 19], [5, 14, 20], [5, 15, 15], [5, 15, 17], [5, 15, 18], [5, 15, 19], [5, 15, 20], [5, 17, 17], [5, 17, 18], [5, 17, 19], [5, 17, 20], [5, 18, 18], [5, 18, 19], [5, 18, 20], [5, 19, 19], [5, 19, 20], [5, 20, 20], [6, 6, 6], [6, 6, 7], [6, 6, 10], [6, 6, 11], [6, 6, 12], [6, 6, 13], [6, 6, 14], [6, 6, 15], [6, 6, 17], [6, 6, 18], [6, 6, 19], [6, 6, 20], [6, 7, 7], [6, 7, 10], [6, 7, 11], [6, 7, 12], [6, 7, 13], [6, 7, 14], [6, 7, 15], [6, 7, 17], [6, 7, 18], [6, 7, 19], [6, 7, 20], [6, 10, 10], [6, 10, 11], [6, 10, 12], [6, 10, 13], [6, 10, 14], [6, 10, 15], [6, 10, 17], [6, 10, 18], [6, 10, 19], [6, 10, 20], [6, 11, 11], [6, 11, 12], [6, 11, 13], [6, 11, 14], [6, 11, 15], [6, 11, 17], [6, 11, 18], [6, 11, 19], [6, 11, 20], [6, 12, 12], [6, 12, 13], [6, 12, 14], [6, 12, 15], [6, 12, 17], [6, 12, 18], [6, 12, 19], [6, 12, 20], [6, 13, 13], [6, 13, 14], [6, 13, 15], [6, 13, 17], [6, 13, 18], [6, 13, 19], [6, 13, 20], [6, 14, 14], [6, 14, 15], [6, 14, 17], [6, 14, 18], [6, 14, 19], [6, 14, 20], [6, 15, 15], [6, 15, 17], [6, 15, 18], [6, 15, 19], [6, 15, 20], [6, 17, 17], [6, 17, 18], [6, 17, 19], [6, 17, 20], [6, 18, 18], [6, 18, 19], [6, 18, 20], [6, 19, 19], [6, 19, 20], [6, 20, 20], [7, 7, 7], [7, 7, 10], [7, 7, 11], [7, 7, 12], [7, 7, 13], [7, 7, 14], [7, 7, 15], [7, 7, 17], [7, 7, 18], [7, 7, 19], [7, 7, 20], [7, 10, 10], [7, 10, 11], [7, 10, 12], [7, 10, 13], [7, 10, 14], [7, 10, 15], [7, 10, 17], [7, 10, 18], [7, 10, 19], [7, 10, 20], [7, 11, 11], [7, 11, 12], [7, 11, 13], [7, 11, 14], [7, 11, 15], [7, 11, 17], [7, 11, 18], [7, 11, 19], [7, 11, 20], [7, 12, 12], [7, 12, 13], [7, 12, 14], [7, 12, 15], [7, 12, 17], [7, 12, 18], [7, 12, 19], [7, 12, 20], [7, 13, 13], [7, 13, 14], [7, 13, 15], [7, 13, 17], [7, 13, 18], [7, 13, 19], [7, 13, 20], [7, 14, 14], [7, 14, 15], [7, 14, 17], [7, 14, 18], [7, 14, 19], [7, 14, 20], [7, 15, 15], [7, 15, 17], [7, 15, 18], [7, 15, 19], [7, 15, 20], [7, 17, 17], [7, 17, 18], [7, 17, 19], [7, 17, 20], [7, 18, 18], [7, 18, 19], [7, 18, 20], [7, 19, 19], [7, 19, 20], [7, 20, 20], [10, 10, 10], [10, 10, 11], [10, 10, 12], [10, 10, 13], [10, 10, 14], [10, 10, 15], [10, 10, 17], [10, 10, 18], [10, 10, 19], [10, 10, 20], [10, 11, 11], [10, 11, 12], [10, 11, 13], [10, 11, 14], [10, 11, 15], [10, 11, 17], [10, 11, 18], [10, 11, 19], [10, 11, 20], [10, 12, 12], [10, 12, 13], [10, 12, 14], [10, 12, 15], [10, 12, 17], [10, 12, 18], [10, 12, 19], [10, 12, 20], [10, 13, 13], [10, 13, 14], [10, 13, 15], [10, 13, 17], [10, 13, 18], [10, 13, 19], [10, 13, 20], [10, 14, 14], [10, 14, 15], [10, 14, 17], [10, 14, 18], [10, 14, 19], [10, 14, 20], [10, 15, 15], [10, 15, 17], [10, 15, 18], [10, 15, 19], [10, 15, 20], [10, 17, 17], [10, 17, 18], [10, 17, 19], [10, 17, 20], [10, 18, 18], [10, 18, 19], [10, 18, 20], [10, 19, 19], [10, 19, 20], [10, 20, 20], [11, 11, 11], [11, 11, 12], [11, 11, 13], [11, 11, 14], [11, 11, 15], [11, 11, 17], [11, 
11, 18], [11, 11, 19], [11, 11, 20], [11, 12, 12], [11, 12, 13], [11, 12, 14], [11, 12, 15], [11, 12, 17], [11, 12, 18], [11, 12, 19], [11, 12, 20], [11, 13, 13], [11, 13, 14], [11, 13, 15], [11, 13, 17], [11, 13, 18], [11, 13, 19], [11, 13, 20], [11, 14, 14], [11, 14, 15], [11, 14, 17], [11, 14, 18], [11, 14, 19], [11, 14, 20], [11, 15, 15], [11, 15, 17], [11, 15, 18], [11, 15, 19], [11, 15, 20], [11, 17, 17], [11, 17, 18], [11, 17, 19], [11, 17, 20], [11, 18, 18], [11, 18, 19], [11, 18, 20], [11, 19, 19], [11, 19, 20], [11, 20, 20], [12, 12, 12], [12, 12, 13], [12, 12, 14], [12, 12, 15], [12, 12, 17], [12, 12, 18], [12, 12, 19], [12, 12, 20], [12, 13, 13], [12, 13, 14], [12, 13, 15], [12, 13, 17], [12, 13, 18], [12, 13, 19], [12, 13, 20], [12, 14, 14], [12, 14, 15], [12, 14, 17], [12, 14, 18], [12, 14, 19], [12, 14, 20], [12, 15, 15], [12, 15, 17], [12, 15, 18], [12, 15, 19], [12, 15, 
20], [12, 17, 17], [12, 17, 18], [12, 17, 19], [12, 17, 20], [12, 18, 18], [12, 18, 19], [12, 18, 20], [12, 19, 19], [12, 19, 20], [12, 20, 20], [13, 13, 13], [13, 13, 14], [13, 13, 15], [13, 13, 17], [13, 13, 18], [13, 13, 19], [13, 13, 20], [13, 14, 14], [13, 14, 15], [13, 14, 17], [13, 14, 18], [13, 14, 19], [13, 14, 20], [13, 15, 15], [13, 15, 17], [13, 15, 18], [13, 15, 19], [13, 15, 20], [13, 17, 17], [13, 17, 18], [13, 17, 19], [13, 17, 20], [13, 18, 18], [13, 18, 19], [13, 18, 20], [13, 19, 19], [13, 19, 20], [13, 20, 20], [14, 14, 14], [14, 14, 15], [14, 14, 17], [14, 14, 18], [14, 14, 19], [14, 14, 20], [14, 15, 15], [14, 15, 17], [14, 15, 18], [14, 15, 19], [14, 15, 20], [14, 17, 17], [14, 17, 18], [14, 17, 19], [14, 17, 20], [14, 18, 18], [14, 18, 19], [14, 18, 20], [14, 19, 19], [14, 19, 20], [14, 20, 20], [15, 15, 15], [15, 15, 17], [15, 15, 18], [15, 15, 19], [15, 15, 20], [15, 17, 17], [15, 17, 18], [15, 17, 19], [15, 17, 20], [15, 18, 18], [15, 18, 19], [15, 18, 20], [15, 19, 19], [15, 19, 20], [15, 20, 20], [17, 17, 17], [17, 17, 18], [17, 17, 19], [17, 17, 20], [17, 18, 18], [17, 18, 19], 
[17, 18, 20], [17, 19, 19], [17, 19, 20], [17, 20, 20], [18, 18, 18], [18, 18, 19], [18, 18, 20], [18, 19, 19], [18, 19, 20], [18, 20, 20], [19, 19, 19], [19, 19, 20], [19, 20, 20], [20, 20, 20]]

with open('difficulties.csv', 'r') as file:
    reader = csv.reader(file)
    difficulties_read = [list(map(float, row)) for row in reader]

# Calculating expected score for board 1 for all dice

def expected_score(board, dice):
    # Initialize variables
    current_board = []
    current_difficulties = []
    expected_score_m1 = 0
    expected_score_m2 = 0
    expected_score_m3 = 0
    dval = extensive_dice_list.index(dice)

    # Populate current_board and current_difficulties from input board
    for val in board:
        current_board.append(val)
        current_difficulties.append(difficulties_read[dval][val])

    # Method 1: Calculate expected score by dividing each board value by its difficulty
    for index, num in enumerate(board):    
        if current_difficulties[index] > -1:    
            expected_score_m1 += num / current_difficulties[index]

    # Method 2: Process the current_board and current_difficulties separately
    # Initialize a time limit
    time_limit = 30
    current_difficulties_2 = current_difficulties.copy()
    current_board_2 = current_board.copy()

    while current_board_2 and current_difficulties_2:
        last_index = len(current_board_2) - 1
        last_difficulty = current_difficulties_2[last_index]

        # Check if the last problem is unsolvable or too difficult
        if last_difficulty == -1 or last_difficulty > 10:
            current_board_2.pop()
            current_difficulties_2.pop()
        # Check if there's not enough time for the last problem
        elif time_limit < last_difficulty:
            current_board_2.pop()
            current_difficulties_2.pop()
        else:
            # Solve the problem and update the score and time limit
            time_limit -= last_difficulty
            expected_score_m2 += current_board_2[last_index]
            current_board_2.pop()
            current_difficulties_2.pop()

    # Method 3
    paired_list = list(zip(board, current_difficulties))
    sorted_by_difficulty = sorted(paired_list, key=lambda x: x[1])  # Sort by difficulty

    time = 30

    for num, difficulty in sorted_by_difficulty:
        if difficulty == -1 or difficulty > 10:
            continue
        elif time < difficulty:
            break
        else:
            time -= difficulty
            expected_score_m3 += num

    multiplier = 39.48/30
    total_score = round(100*(expected_score_m1*0.1 + expected_score_m2*0.7 + expected_score_m3*0.2))*multiplier/100
    return round(100*total_score)/100

def calculate_percentile(data, value):
    """
    Calculate the percentile of a value in a list.

    Args:
        data (list): The list of values.
        value (float): The value for which you want to calculate the percentile.

    Returns:
        float: The percentile of the value in the list.
    """
    # Ensure the data is sorted
    sorted_data = np.sort(data)
    
    # Calculate the percentile using NumPy's percentile function
    percentile = np.percentile(sorted_data, 100 * (np.searchsorted(sorted_data, value) / len(sorted_data)))

    return percentile


def generate_similar_scoring_dice_rolls(board_difficulties, board_index, num_rounds):
    normal_dice_list = [[2, 2, 2], [2, 2, 3], [2, 2, 5], [2, 2, 6], [2, 2, 7], [2, 2, 10], [2, 3, 3], [2, 3, 5], [2, 3, 6],
            [2, 3, 7], [2, 3, 10], [2, 3, 11], [2, 3, 12], [2, 3, 13], [2, 3, 14], [2, 3, 15], [2, 3, 17],
            [2, 3, 18], [2, 3, 19], [2, 3, 20], [2, 5, 5], [2, 5, 6], [2, 5, 10], [2, 6, 6], [2, 6, 7], [2, 6, 12],
            [3, 3, 5], [3, 3, 6], [3, 3, 12], [3, 5, 5], [3, 5, 6], [3, 5, 7], [3, 5, 15], [3, 6, 6],
            [3, 6, 7], [3, 6, 10], [3, 6, 12], [3, 6, 18]]
    
    extensive_dice_list = [[1, 2, 2], [1, 2, 3], [1, 2, 5], [1, 2, 6], [1, 2, 7], [1, 2, 10], [1, 2, 11], [1, 2, 12], [1, 2, 13], [1, 2, 14], [1, 2, 15], [1, 2, 17], [1, 2, 18], [1, 2, 19], [1, 2, 20], [1, 3, 3], [1, 3, 5], [1, 3, 6], [1, 3, 7], [1, 3, 10], [1, 3, 11], [1, 3, 12], [1, 3, 13], [1, 3, 14], [1, 3, 15], [1, 3, 17], [1, 3, 18], [1, 3, 19], [1, 3, 20], [1, 5, 5], [1, 5, 6], [1, 5, 7], [1, 5, 10], [1, 5, 11], [1, 5, 12], [1, 5, 13], [1, 5, 14], [1, 5, 15], [1, 5, 17], [1, 5, 18], [1, 5, 19], [1, 5, 20], [1, 6, 6], [1, 6, 7], [1, 6, 10], [1, 6, 11], [1, 6, 12], [1, 6, 13], [1, 6, 14], [1, 6, 15], [1, 6, 17], [1, 6, 18], [1, 6, 19], [1, 6, 20], [1, 7, 7], [1, 7, 10], [1, 7, 11], [1, 7, 12], [1, 7, 13], [1, 7, 14], [1, 7, 15], [1, 7, 17], [1, 7, 18], [1, 7, 19], [1, 7, 20], [1, 10, 10], [1, 10, 11], [1, 10, 12], [1, 10, 13], [1, 10, 14], [1, 10, 15], [1, 10, 17], [1, 10, 18], [1, 10, 19], [1, 10, 20], [1, 
11, 11], [1, 11, 12], [1, 11, 13], [1, 11, 14], [1, 11, 15], [1, 11, 17], [1, 11, 18], [1, 11, 19], [1, 11, 20], [1, 12, 12], [1, 12, 13], [1, 12, 14], [1, 12, 15], [1, 12, 17], [1, 12, 18], [1, 12, 19], [1, 12, 20], [1, 13, 
13], [1, 13, 14], [1, 13, 15], [1, 13, 17], [1, 13, 18], [1, 13, 19], [1, 13, 20], [1, 14, 14], [1, 14, 15], [1, 14, 17], [1, 14, 18], [1, 14, 19], [1, 14, 20], [1, 15, 15], [1, 15, 17], [1, 15, 18], [1, 15, 19], [1, 15, 20], [1, 17, 17], [1, 17, 18], [1, 17, 19], [1, 17, 20], [1, 18, 18], [1, 18, 19], [1, 18, 20], [1, 19, 19], [1, 19, 20], [1, 20, 20], [2, 2, 2], [2, 2, 3], [2, 2, 5], [2, 2, 6], [2, 2, 7], [2, 2, 10], [2, 2, 11], [2, 2, 12], [2, 2, 13], [2, 2, 14], [2, 2, 15], [2, 2, 17], [2, 2, 18], [2, 2, 19], [2, 2, 20], [2, 3, 3], [2, 3, 5], [2, 3, 6], [2, 3, 7], [2, 3, 10], [2, 3, 11], [2, 3, 12], [2, 3, 13], [2, 3, 14], [2, 3, 15], [2, 3, 17], [2, 3, 18], [2, 
3, 19], [2, 3, 20], [2, 5, 5], [2, 5, 6], [2, 5, 7], [2, 5, 10], [2, 5, 11], [2, 5, 12], [2, 5, 13], [2, 5, 14], [2, 5, 15], [2, 5, 17], [2, 5, 18], [2, 5, 19], [2, 5, 20], [2, 6, 6], [2, 6, 7], [2, 6, 10], [2, 6, 11], [2, 6, 12], [2, 6, 13], [2, 6, 14], [2, 6, 15], [2, 6, 17], [2, 6, 18], [2, 6, 19], [2, 6, 20], [2, 7, 7], [2, 7, 10], [2, 7, 11], [2, 7, 12], [2, 7, 13], [2, 7, 14], [2, 7, 15], [2, 7, 17], [2, 7, 18], [2, 7, 19], [2, 7, 20], [2, 
10, 10], [2, 10, 11], [2, 10, 12], [2, 10, 13], [2, 10, 14], [2, 10, 15], [2, 10, 17], [2, 10, 18], [2, 10, 19], [2, 10, 20], [2, 11, 11], [2, 11, 12], [2, 11, 13], [2, 11, 14], [2, 11, 15], [2, 11, 17], [2, 11, 18], [2, 11, 
19], [2, 11, 20], [2, 12, 12], [2, 12, 13], [2, 12, 14], [2, 12, 15], [2, 12, 17], [2, 12, 18], [2, 12, 19], [2, 12, 20], [2, 13, 13], [2, 13, 14], [2, 13, 15], [2, 13, 17], [2, 13, 18], [2, 13, 19], [2, 13, 20], [2, 14, 14], [2, 14, 15], [2, 14, 17], [2, 14, 18], [2, 14, 19], [2, 14, 20], [2, 15, 15], [2, 15, 17], [2, 15, 18], [2, 15, 19], [2, 15, 20], [2, 17, 17], [2, 17, 18], [2, 17, 19], [2, 17, 20], [2, 18, 18], [2, 18, 19], [2, 18, 20], [2, 19, 19], [2, 19, 20], [2, 20, 20], [3, 3, 3], [3, 3, 5], [3, 3, 6], [3, 3, 7], [3, 3, 10], [3, 3, 11], [3, 3, 12], [3, 3, 13], [3, 3, 14], [3, 3, 15], [3, 3, 17], [3, 3, 18], [3, 3, 19], [3, 3, 20], [3, 5, 5], [3, 5, 6], [3, 5, 7], [3, 5, 10], [3, 5, 11], [3, 5, 12], [3, 5, 13], [3, 5, 14], [3, 5, 15], [3, 5, 17], [3, 5, 18], [3, 5, 19], [3, 5, 20], [3, 6, 6], [3, 6, 7], [3, 6, 10], [3, 6, 11], [3, 6, 12], [3, 6, 13], [3, 6, 14], [3, 6, 15], [3, 6, 17], [3, 6, 18], [3, 6, 19], [3, 6, 20], [3, 7, 7], [3, 7, 10], [3, 7, 11], [3, 7, 12], [3, 7, 13], [3, 7, 14], [3, 7, 15], [3, 7, 17], [3, 7, 18], [3, 7, 19], [3, 7, 20], [3, 10, 10], [3, 10, 11], [3, 10, 12], [3, 10, 13], [3, 10, 14], [3, 10, 15], [3, 10, 17], [3, 10, 18], [3, 10, 19], [3, 10, 20], [3, 11, 11], [3, 11, 12], [3, 11, 13], [3, 11, 14], [3, 11, 15], [3, 11, 17], [3, 11, 18], [3, 11, 19], [3, 11, 20], [3, 12, 12], [3, 12, 13], [3, 12, 14], [3, 12, 15], [3, 12, 17], [3, 12, 18], [3, 12, 19], [3, 12, 20], [3, 13, 13], [3, 13, 14], [3, 13, 15], [3, 13, 17], [3, 13, 18], [3, 13, 19], [3, 13, 20], [3, 14, 14], [3, 14, 15], [3, 14, 17], [3, 14, 18], [3, 14, 19], [3, 14, 20], [3, 15, 15], [3, 15, 17], [3, 15, 18], [3, 15, 19], [3, 15, 20], [3, 17, 17], [3, 17, 18], [3, 17, 19], [3, 17, 20], [3, 18, 18], [3, 18, 19], [3, 18, 20], [3, 19, 19], [3, 19, 20], [3, 20, 20], [5, 5, 5], [5, 5, 6], [5, 5, 7], [5, 5, 10], [5, 5, 11], [5, 5, 12], [5, 5, 13], [5, 5, 14], [5, 5, 15], [5, 5, 17], [5, 5, 18], [5, 5, 19], [5, 5, 20], [5, 6, 6], [5, 6, 7], [5, 6, 10], [5, 6, 11], [5, 6, 12], [5, 6, 13], [5, 6, 14], [5, 6, 15], [5, 6, 17], [5, 6, 18], [5, 6, 19], [5, 6, 20], [5, 7, 7], [5, 7, 10], [5, 7, 11], [5, 7, 12], [5, 7, 13], [5, 7, 14], [5, 7, 15], [5, 7, 17], [5, 7, 18], [5, 7, 19], [5, 7, 20], [5, 10, 10], [5, 10, 11], [5, 10, 12], [5, 10, 13], [5, 10, 14], [5, 10, 15], [5, 10, 17], [5, 10, 18], [5, 10, 19], [5, 10, 20], [5, 11, 11], [5, 11, 12], [5, 11, 13], [5, 11, 14], [5, 11, 15], [5, 11, 17], [5, 11, 18], [5, 11, 19], [5, 11, 20], [5, 12, 12], [5, 12, 13], [5, 12, 14], [5, 12, 15], [5, 12, 17], [5, 12, 18], [5, 12, 19], [5, 12, 20], [5, 13, 13], [5, 13, 14], [5, 13, 15], [5, 13, 17], [5, 13, 18], [5, 13, 19], [5, 13, 20], [5, 14, 14], [5, 14, 15], [5, 14, 17], [5, 14, 18], [5, 14, 19], [5, 14, 20], [5, 15, 15], [5, 15, 17], [5, 15, 18], [5, 15, 19], [5, 15, 20], [5, 17, 17], [5, 17, 18], [5, 17, 19], [5, 17, 20], [5, 18, 18], [5, 18, 19], [5, 18, 20], [5, 19, 19], [5, 19, 20], [5, 20, 20], [6, 6, 6], [6, 6, 7], [6, 6, 10], [6, 6, 11], [6, 6, 12], [6, 6, 13], [6, 6, 14], [6, 6, 15], [6, 6, 17], [6, 6, 18], [6, 6, 19], [6, 6, 20], [6, 7, 7], [6, 7, 10], [6, 7, 11], [6, 7, 12], [6, 7, 13], [6, 7, 14], [6, 7, 15], [6, 7, 17], [6, 7, 18], [6, 7, 19], [6, 7, 20], [6, 10, 10], [6, 10, 11], [6, 10, 12], [6, 10, 13], [6, 10, 14], [6, 10, 15], [6, 10, 17], [6, 10, 18], [6, 10, 19], [6, 10, 20], [6, 11, 11], [6, 11, 12], [6, 11, 13], [6, 11, 14], [6, 11, 15], [6, 11, 17], [6, 11, 18], [6, 11, 19], [6, 11, 20], [6, 12, 12], [6, 12, 13], [6, 12, 14], [6, 12, 15], [6, 12, 17], [6, 12, 18], [6, 12, 19], [6, 12, 20], [6, 13, 13], [6, 13, 14], [6, 13, 15], [6, 13, 17], [6, 13, 18], [6, 13, 19], [6, 13, 20], [6, 14, 14], [6, 14, 15], [6, 14, 17], [6, 14, 18], [6, 14, 19], [6, 14, 20], [6, 15, 15], [6, 15, 17], [6, 15, 18], [6, 15, 19], [6, 15, 20], [6, 17, 17], [6, 17, 18], [6, 17, 19], [6, 17, 20], [6, 18, 18], [6, 18, 19], [6, 18, 20], [6, 19, 19], [6, 19, 20], [6, 20, 20], [7, 7, 7], [7, 7, 10], [7, 7, 11], [7, 7, 12], [7, 7, 13], [7, 7, 14], [7, 7, 15], [7, 7, 17], [7, 7, 18], [7, 7, 19], [7, 7, 20], [7, 10, 10], [7, 10, 11], [7, 10, 12], [7, 10, 13], [7, 10, 14], [7, 10, 15], [7, 10, 17], [7, 10, 18], [7, 10, 19], [7, 10, 20], [7, 11, 11], [7, 11, 12], [7, 11, 13], [7, 11, 14], [7, 11, 15], [7, 11, 17], [7, 11, 18], [7, 11, 19], [7, 11, 20], [7, 12, 12], [7, 12, 13], [7, 12, 14], [7, 12, 15], [7, 12, 17], [7, 12, 18], [7, 12, 19], [7, 12, 20], [7, 13, 13], [7, 13, 14], [7, 13, 15], [7, 13, 17], [7, 13, 18], [7, 13, 19], [7, 13, 20], [7, 14, 14], [7, 14, 15], [7, 14, 17], [7, 14, 18], [7, 14, 19], [7, 14, 20], [7, 15, 15], [7, 15, 17], [7, 15, 18], [7, 15, 19], [7, 15, 20], [7, 17, 17], [7, 17, 18], [7, 17, 19], [7, 17, 20], [7, 18, 18], [7, 18, 19], [7, 18, 20], [7, 19, 19], [7, 19, 20], [7, 20, 20], [10, 10, 10], [10, 10, 11], [10, 10, 12], [10, 10, 13], [10, 10, 14], [10, 10, 15], [10, 10, 17], [10, 10, 18], [10, 10, 19], [10, 10, 20], [10, 11, 11], [10, 11, 12], [10, 11, 13], [10, 11, 14], [10, 11, 15], [10, 11, 17], [10, 11, 18], [10, 11, 19], [10, 11, 20], [10, 12, 12], [10, 12, 13], [10, 12, 14], [10, 12, 15], [10, 12, 17], [10, 12, 18], [10, 12, 19], [10, 12, 20], [10, 13, 13], [10, 13, 14], [10, 13, 15], [10, 13, 17], [10, 13, 18], [10, 13, 19], [10, 13, 20], [10, 14, 14], [10, 14, 15], [10, 14, 17], [10, 14, 18], [10, 14, 19], [10, 14, 20], [10, 15, 15], [10, 15, 17], [10, 15, 18], [10, 15, 19], [10, 15, 20], [10, 17, 17], [10, 17, 18], [10, 17, 19], [10, 17, 20], [10, 18, 18], [10, 18, 19], [10, 18, 20], [10, 19, 19], [10, 19, 20], [10, 20, 20], [11, 11, 11], [11, 11, 12], [11, 11, 13], [11, 11, 14], [11, 11, 15], [11, 11, 17], [11, 
11, 18], [11, 11, 19], [11, 11, 20], [11, 12, 12], [11, 12, 13], [11, 12, 14], [11, 12, 15], [11, 12, 17], [11, 12, 18], [11, 12, 19], [11, 12, 20], [11, 13, 13], [11, 13, 14], [11, 13, 15], [11, 13, 17], [11, 13, 18], [11, 13, 19], [11, 13, 20], [11, 14, 14], [11, 14, 15], [11, 14, 17], [11, 14, 18], [11, 14, 19], [11, 14, 20], [11, 15, 15], [11, 15, 17], [11, 15, 18], [11, 15, 19], [11, 15, 20], [11, 17, 17], [11, 17, 18], [11, 17, 19], [11, 17, 20], [11, 18, 18], [11, 18, 19], [11, 18, 20], [11, 19, 19], [11, 19, 20], [11, 20, 20], [12, 12, 12], [12, 12, 13], [12, 12, 14], [12, 12, 15], [12, 12, 17], [12, 12, 18], [12, 12, 19], [12, 12, 20], [12, 13, 13], [12, 13, 14], [12, 13, 15], [12, 13, 17], [12, 13, 18], [12, 13, 19], [12, 13, 20], [12, 14, 14], [12, 14, 15], [12, 14, 17], [12, 14, 18], [12, 14, 19], [12, 14, 20], [12, 15, 15], [12, 15, 17], [12, 15, 18], [12, 15, 19], [12, 15, 
20], [12, 17, 17], [12, 17, 18], [12, 17, 19], [12, 17, 20], [12, 18, 18], [12, 18, 19], [12, 18, 20], [12, 19, 19], [12, 19, 20], [12, 20, 20], [13, 13, 13], [13, 13, 14], [13, 13, 15], [13, 13, 17], [13, 13, 18], [13, 13, 19], [13, 13, 20], [13, 14, 14], [13, 14, 15], [13, 14, 17], [13, 14, 18], [13, 14, 19], [13, 14, 20], [13, 15, 15], [13, 15, 17], [13, 15, 18], [13, 15, 19], [13, 15, 20], [13, 17, 17], [13, 17, 18], [13, 17, 19], [13, 17, 20], [13, 18, 18], [13, 18, 19], [13, 18, 20], [13, 19, 19], [13, 19, 20], [13, 20, 20], [14, 14, 14], [14, 14, 15], [14, 14, 17], [14, 14, 18], [14, 14, 19], [14, 14, 20], [14, 15, 15], [14, 15, 17], [14, 15, 18], [14, 15, 19], [14, 15, 20], [14, 17, 17], [14, 17, 18], [14, 17, 19], [14, 17, 20], [14, 18, 18], [14, 18, 19], [14, 18, 20], [14, 19, 19], [14, 19, 20], [14, 20, 20], [15, 15, 15], [15, 15, 17], [15, 15, 18], [15, 15, 19], [15, 15, 20], [15, 17, 17], [15, 17, 18], [15, 17, 19], [15, 17, 20], [15, 18, 18], [15, 18, 19], [15, 18, 20], [15, 19, 19], [15, 19, 20], [15, 20, 20], [17, 17, 17], [17, 17, 18], [17, 17, 19], [17, 17, 20], [17, 18, 18], [17, 18, 19], 
[17, 18, 20], [17, 19, 19], [17, 19, 20], [17, 20, 20], [18, 18, 18], [18, 18, 19], [18, 18, 20], [18, 19, 19], [18, 19, 20], [18, 20, 20], [19, 19, 19], [19, 19, 20], [19, 20, 20], [20, 20, 20]]

    selected_board = board_difficulties[board_index]
    sorted_scores = sorted(enumerate(selected_board), key=lambda x: x[1], reverse=True)

    # Percentile thresholds
    top_5_percent_index = int(len(selected_board) * 0.05)
    top_10_percent_index = int(len(selected_board) * 0.10)
    top_25_percent_index = int(len(selected_board) * 0.30)
    top_50_percent_index = int(len(selected_board) * 0.70)

    # Player scores
    p1_scores = []
    p2_scores = []

    # Percentile lists for data for competition generation summary
    p1_percentiles = []
    p2_percentiles = []


    dice_roll_pairs = []
    selected_pairs = set()  # To keep track of selected pairs and avoid duplicates

    while len(dice_roll_pairs) < num_rounds:
        # Decide which score bracket to choose from
        bracket_choice = random.choices(['top_5', 'top_10', 'top_25', 'top_50'], weights=[8, 5, 3, 1], k=1)[0]

        # Determine the range for random selection based on the chosen bracket
        if bracket_choice == 'top_5':
            range_start, range_end = 0, top_5_percent_index
        elif bracket_choice == 'top_10':
            range_start, range_end = 0, top_10_percent_index
        elif bracket_choice == 'top_25':
            range_start, range_end = 0, top_25_percent_index
        else:  # 'top_50'
            range_start, range_end = 0, top_50_percent_index

        # Weighted random choice of starting position within the chosen bracket
        pos = random.randint(range_start, range_end - 1)

        # If we are on an even number of dice, give p1 the better score. else, give it two p2

        if len(dice_roll_pairs) % 2 == 0:
            first_roll_index, _score = sorted_scores[pos]
            p1_scores.append(round(_score))
            p1_percentiles.append(100 * pos / len(sorted_scores))
            second_roll_index, _score = sorted_scores[pos + 1]
            p2_percentiles.append(100 * (pos+1) / len(sorted_scores))
            p2_scores.append(round(_score))
            pair = (tuple(extensive_dice_list[first_roll_index]), tuple(extensive_dice_list[second_roll_index]))
        else: 
            # Same code as previous, but players are swtiched
            second_roll_index, _score = sorted_scores[pos]
            p2_scores.append(round(_score))
            p2_percentiles.append(100 * (pos+1) / len(sorted_scores))
            first_roll_index, _score = sorted_scores[pos + 1]
            p1_scores.append(round(_score))
            p1_percentiles.append(100 * pos / len(sorted_scores))
            pair = (tuple(extensive_dice_list[first_roll_index]), tuple(extensive_dice_list[second_roll_index]))


        if pair not in selected_pairs:
            dice_roll_pairs.append(pair)
            selected_pairs.add(pair)

    scores = [p1_scores, p2_scores]
    percentiles = [p1_percentiles, p2_percentiles]

    return [scores, dice_roll_pairs, percentiles]


board_difficulties = []

for index, board in enumerate(boards):
    board_difficulties.append([])
    for dice in extensive_dice_list:
        board_difficulties[index].append(expected_score(board, dice))

boards_dice = []
scores = []
percentiles = []

for i in range(len(boards)):
    if i == len(boards)-1:
        generated = generate_similar_scoring_dice_rolls(board_difficulties, i, num_rounds=dice_count+extra_dice)
        boards_dice.append(generated[1])
        scores.append(generated[0])
        percentiles.append(generated[2])
    else:
        generated = generate_similar_scoring_dice_rolls(board_difficulties, i, num_rounds=dice_count)
        boards_dice.append(generated[1])
        scores.append(generated[0])
        percentiles.append(generated[2])


def possible_score(board, dice):
    possible_score = 0
    for num in board:
        if difficulties_read[extensive_dice_list.index(dice)][num] != -1:
            possible_score += num

    return possible_score


def convert_dice_data(dice_data):
    rounds = dice_data[0]
     # Initialize the result list
    converted_data = []

    for round in rounds:
        # Unpack the tuples for each player
        roll_p1, roll_p2 = round

        # Convert tuples to lists and add them to the result
        converted_data.append([list(roll_p1), list(roll_p2)])

    return converted_data

converted_dice = convert_dice_data(boards_dice)



# Doc generation
import docx
from docx.shared import Pt




# Define functions to work with python-docx
def create_doc():
    """Create a new Word document and return it."""
    doc = docx.Document()
    return doc

def add_table_to_doc(doc, board):
    """Add a 6x6 table to the document for a given board."""
    table = doc.add_table(rows=6, cols=6)
    for row in range(6):
        for col in range(6):
            table.cell(row, col).text = str(board[row * 6 + col])
            # Set the font size
            for paragraph in table.cell(row, col).paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(28)
    return doc


def add_dice_rolls_to_doc(doc, dice_rolls_data):
    """Add a list of dice rolls to the document in a formatted way with specific formatting."""
    for round_number, (roll_p1, roll_p2) in enumerate(dice_rolls_data, start=1):
        p = doc.add_paragraph()

        # Set the font size for the paragraph
        run = p.add_run(f"Round {round_number}\n")
        font = run.font
        font.size = Pt(20)
        font.bold = True

        # Format the dice rolls for each player and add them to the paragraph
        roll_p1_formatted = ", ".join(map(str, roll_p1))
        roll_p2_formatted = ", ".join(map(str, roll_p2))

        run = p.add_run(f"Player 1: {roll_p1_formatted}                    Player 2: {roll_p2_formatted}\n")
        run.font.size = Pt(20)  # Set the font size for each player's roll
        
    return doc







# Generate Word documents for each board
for i, board in enumerate(boards):
    doc = create_doc()
    doc = add_table_to_doc(doc, board)

    # Assuming dice_rolls is a list of tuples representing dice rolls for each board
    # This part of the code needs to be adapted based on how dice_rolls are stored in your program
    dice_rolls_for_board = boards_dice[i]  # Replace with actual dice rolls for the current board
    doc = add_dice_rolls_to_doc(doc, dice_rolls_for_board)

    doc.save(f'board_{i+1}.docx')

# Generate a summary page
def add_heading(doc, text, level=1):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)
    return doc

def add_paragraph(doc, text, font_size=12, bold=False, alignment=None):
    """Add a paragraph to the document with specific formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment:
        p.alignment = alignment
    return doc


def add_table(doc, data):
    """Add a table to the document with the given data."""
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.autofit = False
    table.allow_autofit = False

    for i, row in enumerate(data):
        for j, cell_value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(cell_value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc

# Create the summary document
summary_doc = create_doc()

# Title and introduction
summary_doc = add_heading(summary_doc, 'Competition Generation Summary', level=1)
summary_doc = add_paragraph(summary_doc, 'This document summarizes the setup and expected outcomes of the generated competition.', font_size=14, bold=True)

# Board and Dice Roll Overview
summary_doc = add_heading(summary_doc, 'Board and Dice Roll Overview', level=2)
summary_doc = add_paragraph(summary_doc, 'For this competition, a total of ' + str(board_count) + ' game boards were generated, for a total of ' + str(round_count) + ' rounds', font_size=12)

board_mins = []
board_maxs = []

for index, board in enumerate(boards):
    board_mins.append(min(board))
    board_maxs.append(max(board))

board_ranges_string = ""

for index, board in enumerate(boards):
    board_ranges_string += "(" + str(board_mins[index]) + "-" + str(board_maxs[index]) + ")   "

summary_doc = add_paragraph(summary_doc, 'The board ranges were the following: ' + board_ranges_string, font_size=12)


# Expected Scores Analysis
summary_doc = add_heading(summary_doc, 'Expected Scores Analysis', level=2)
summary_doc = add_paragraph(summary_doc, 'Based on the board layouts and dice rolls, the expected scores for each player were calculated. Here are the anticipated scores:', font_size=12)


string_scores = ""

for index, board in enumerate(scores):
    if index == 0:
        string_scores += ("Board " + str(index+1))
    else:
        string_scores += ("\nBoard " + str(index+1))
    for p_index, player in enumerate(board):
        string_scores += ("\nPlayer " + str(p_index+1) + ": ")
        for score in player:
            string_scores += (str(score) + "   ")        


summary_doc =add_paragraph(summary_doc, string_scores, font_size=12)

# Difficulty Level Insights
summary_doc = add_heading(summary_doc, 'Dice Roll Percentiles', level=2)

string_percentiles = ""

for index, board in enumerate(percentiles):
    if index == 0:
        string_percentiles += ("Board " + str(index+1))
    else:
        string_percentiles += ("\nBoard " + str(index+1))
    for p_index, player in enumerate(board):
        string_percentiles += ("\nPlayer " + str(p_index+1) + ": ")
        for percentile in player:
            string_percentiles += (str(round(10*percentile)/10) + "%    ")        

summary_doc =add_paragraph(summary_doc, "The dice roll percentile indicates the dice roll's difficulty level for a given board. A lower percentile corresponds to lower difficulty.", font_size=12)
summary_doc =add_paragraph(summary_doc, string_percentiles, font_size=12)

summary_doc = add_heading(summary_doc, '\nPossibility Percentages', level=2)
summary_doc =add_paragraph(summary_doc, "The following figures refer to what percent of the board is possible for a player knock off.", font_size=12)

string_possibles = ""

p1_working_string = ""
p2_working_string = ""

for index, board in enumerate(boards_dice):
    if index == 0:
        string_possibles += ("Board " + str(index+1))
    else:
        string_possibles += ("\nBoard " + str(index+1))
    for _round in board:
        for p_index, player in enumerate(_round):
            if p_index == 0:
                # Player 1
                pos_score = possible_score(boards[index], list(player))
                total_impossible = sum_list(boards[index])
                p1_working_string += str( round(1000*pos_score/total_impossible)/10)  + "%   "
            else:
                pos_score = possible_score(boards[index], list(player))
                total_impossible = sum_list(boards[index])
                p2_working_string += str( round(1000*pos_score/total_impossible)/10)  + "%   "
    string_possibles += "\nPlayer 1: " + p1_working_string
    string_possibles += "\nPlayer 2: " + p2_working_string
    p1_working_string = ""
    p2_working_string = ""    
        

summary_doc =add_paragraph(summary_doc, string_possibles, font_size=12)



# Save the document
summary_doc.save('competition_generation_summary.docx')

print("Competition Generated")

# To do: Make normalized dice more common than extensive,